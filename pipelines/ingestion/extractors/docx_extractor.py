import io
import logging
from typing import List
import docx

from pipelines.ingestion.extractors.extracted_page import ExtractedPage

logger = logging.getLogger("docx-extractor")

def extract_docx(file_bytes: bytes) -> List[ExtractedPage]:
    """
    Extracts text and tables from DOCX bytes.
    - Extracts paragraphs, tables, and headers separately.
    - Tracks heading levels (Heading 1, etc.) to preserve structural metadata.
    """
    extracted_pages = []
    
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"Failed to open DOCX with python-docx: {e}")
        return []

    page_num = 1
    current_section = "General"
    current_level = "h0"
    
    # 1. Extract Document Headers (running headers in page layouts)
    header_texts = []
    for section in doc.sections:
        if section.header and section.header.text.strip():
            header_texts.append(section.header.text.strip())
            
    if header_texts:
        extracted_pages.append(ExtractedPage(
            page_number=page_num,
            content="\n".join(header_texts),
            content_type="text",
            metadata={"type": "header", "page_number": page_num}
        ))
        page_num += 1

    # 2. Extract Paragraphs & track heading hierarchy
    paragraph_texts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Detect if paragraph is a heading (e.g. style name contains 'Heading')
        style_name = p.style.name if p.style else ""
        if "Heading" in style_name:
            # e.g., "Heading 1" -> h1
            try:
                level_num = style_name.split()[-1]
                current_level = f"h{level_num}"
            except Exception:
                current_level = "h1"
            current_section = text
            
        paragraph_texts.append(text)
        
    if paragraph_texts:
        extracted_pages.append(ExtractedPage(
            page_number=page_num,
            content="\n\n".join(paragraph_texts),
            content_type="text",
            metadata={
                "page_number": page_num,
                "level": current_level,
                "section": current_section
            }
        ))
        page_num += 1

    # 3. Extract Tables and represent them as Markdown tables
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
            
        # Convert row cells list to markdown
        if table_rows:
            col_count = len(table_rows[0])
            markdown_lines = []
            
            # Header Row
            markdown_lines.append("| " + " | ".join(table_rows[0]) + " |")
            # Separator
            markdown_lines.append("| " + " | ".join(["---"] * col_count) + " |")
            # Data Rows
            for row in table_rows[1:]:
                padded_row = row[:col_count] + [""] * (col_count - len(row))
                markdown_lines.append("| " + " | ".join(padded_row) + " |")
                
            markdown_table = "\n".join(markdown_lines)
            if markdown_table.strip():
                extracted_pages.append(ExtractedPage(
                    page_number=page_num,
                    content=markdown_table,
                    content_type="table",
                    metadata={
                        "page_number": page_num,
                        "table_index": table_idx,
                        "section": current_section
                    }
                ))
                page_num += 1

    return extracted_pages
